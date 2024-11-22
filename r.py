from flask import Flask, render_template, request, redirect, url_for
import os
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

# إعداد مجلدات التخزين
UPLOAD_FOLDER = 'uploads'
REPORTS_FOLDER = 'reports'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# الصفحة الرئيسية
@app.route('/')
def index():
    return render_template('index.html')

# رفع التقرير وإنشاء ملف Word
@app.route('/submit', methods=['POST'])
def submit_report():
    try:
        # استلام بيانات التقرير
        case_id = request.form['case_id']
        field = request.form['field']
        description = request.form['description']
        images = request.files.getlist('images')

        # حفظ الصور
        image_paths = []
        for image in images:
            if image.filename:  # تحقق إذا تم رفع صورة
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], image.filename)
                image.save(image_path)
                image_paths.append(image_path)

        # إنشاء ملف Word
        output_path = os.path.join(REPORTS_FOLDER, f"{case_id}.docx")
        create_word_report(case_id, field, description, image_paths, output_path)

        # إعادة التوجيه إلى صفحة نجاح
        return redirect(url_for('report_success', case_id=case_id))
    except Exception as e:
        return f"Error occurred: {e}", 500

# وظيفة لإنشاء تقرير Word
def create_word_report(case_id, field, description, image_paths, output_path):
    doc = Document()
    doc.add_heading('Maintenance Report', level=1)

    # إضافة الحقول الأساسية
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_paragraph(f"Field: {field}")
    doc.add_paragraph(f"Description: {description}")

    # إضافة الصور (إن وجدت)
    if image_paths:
        doc.add_heading('Attached Images', level=2)
        for image_path in image_paths:
            doc.add_picture(image_path, width=Inches(2))

    doc.save(output_path)

# صفحة نجاح التقرير
@app.route('/success/<case_id>')
def report_success(case_id):
    report_path = os.path.join(REPORTS_FOLDER, f"{case_id}.docx")
    return render_template('success.html', report_path=report_path, case_id=case_id)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)  # تشغيل الخادم على الشبكة المحلية
